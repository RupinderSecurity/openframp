#!/usr/bin/env python3
"""
OpenFRAMP SSP Parser — Extracts structured data from FedRAMP SSP documents.

Handles two document types:
1. Main SSP template: system info, contacts, leveraged services, ports/protocols
2. Appendix A: control implementation statements (323 Moderate controls)

Usage:
    python3 ssp_parser.py <path-to-ssp.docx> [--output results.json]
    python3 ssp_parser.py <main-ssp.docx> --appendix-a <appendix-a.docx> [--output results.json]
"""

import json
import re
import sys
from docx import Document


# ── CONTROL PARSER (Appendix A) ──

CONTROL_ID_PATTERN = re.compile(r'^([A-Z]{2}-\d+(?:\(\d+\))?)\s')


def parse_summary_table(table):
    summary = {
        "control_id": "",
        "responsible_role": "",
        "parameters": {},
        "implementation_status": "",
        "control_origination": ""
    }
    
    for row in table.rows:
        cell_text = row.cells[0].text.strip()
        
        if "Control Summary Information" in cell_text:
            match = CONTROL_ID_PATTERN.match(cell_text)
            if match:
                summary["control_id"] = match.group(1)
        
        elif cell_text.startswith("Responsible Role"):
            if len(row.cells) > 1:
                summary["responsible_role"] = row.cells[1].text.strip()
            else:
                parts = cell_text.split(":", 1)
                if len(parts) > 1:
                    summary["responsible_role"] = parts[1].strip()
        
        elif cell_text.startswith("Parameter"):
            param_match = re.match(r'Parameter\s+([^:]+):', cell_text)
            if param_match:
                param_name = param_match.group(1).strip()
                if len(row.cells) > 1:
                    value = row.cells[1].text.strip()
                else:
                    parts = cell_text.split(":", 1)
                    value = parts[1].strip() if len(parts) > 1 else ""
                summary["parameters"][param_name] = value
        
        elif "Implementation Status" in cell_text:
            summary["implementation_status"] = cell_text
        
        elif "Control Origination" in cell_text:
            summary["control_origination"] = cell_text
    
    return summary


def parse_implementation_table(table):
    implementation = {"control_id": "", "parts": {}}
    
    for row in table.rows:
        cell_text = row.cells[0].text.strip()
        
        if "What is the solution" in cell_text:
            match = CONTROL_ID_PATTERN.match(cell_text)
            if match:
                implementation["control_id"] = match.group(1)
        
        elif cell_text.startswith("Part "):
            part_match = re.match(r'Part\s+([a-z]):', cell_text)
            if part_match:
                part_letter = part_match.group(1)
                if len(row.cells) > 1:
                    narrative = row.cells[1].text.strip()
                else:
                    parts = cell_text.split(":", 1)
                    narrative = parts[1].strip() if len(parts) > 1 else ""
                implementation["parts"][part_letter] = narrative
    
    return implementation


def parse_appendix_a(docx_path):
    print(f"  Loading Appendix A: {docx_path}...")
    doc = Document(docx_path)
    print(f"  Found {len(doc.tables)} tables")
    
    controls = {}
    
    for table in doc.tables:
        first_cell = table.rows[0].cells[0].text.strip() if table.rows else ""
        
        if "Control Summary Information" in first_cell:
            summary = parse_summary_table(table)
            cid = summary["control_id"]
            if cid:
                if cid not in controls:
                    controls[cid] = {"summary": {}, "implementation": {}}
                controls[cid]["summary"] = summary
        
        elif "What is the solution" in first_cell:
            impl = parse_implementation_table(table)
            cid = impl["control_id"]
            if cid:
                if cid not in controls:
                    controls[cid] = {"summary": {}, "implementation": {}}
                controls[cid]["implementation"] = impl
    
    return controls


# ── SYSTEM INFO PARSER (Main SSP) ──

def parse_system_information(table):
    """Extract system characteristics from the System Information table."""
    info = {}
    field_map = {
        "CSP Name:": "csp_name",
        "CSO Name:": "cso_name",
        "FedRAMP Package ID:": "fedramp_package_id",
        "Service Model:": "service_model",
        "FIPS PUB 199 Level": "fips_199_level",
        "Fully Operational as of:": "operational_date",
        "Deployment Model:": "deployment_model",
        "Authorization Path:": "authorization_path",
        "General System Description:": "description",
        "Digital Identity Level": "digital_identity_level"
    }
    
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip()
        value = row.cells[1].text.strip()
        
        for key_prefix, field_name in field_map.items():
            if label.startswith(key_prefix):
                # Strip placeholder text
                if value.startswith("<") and value.endswith(">"):
                    value = ""
                elif "<Insert" in value or "<Choose" in value or "<Enter" in value:
                    value = ""
                info[field_name] = value
                break
    
    return info


def parse_contact_table(table):
    """Extract point of contact from a contact table (System Owner, ISSO, etc.)."""
    contact = {}
    field_map = {
        "Name": "name",
        "Title": "title",
        "Company / Organization": "organization",
        "Company/Organization": "organization",
        "Address": "address",
        "Phone Number": "phone",
        "Phone": "phone",
        "Email Address": "email",
        "Email": "email"
    }
    
    # Get the contact type from the header row
    header = table.rows[0].cells[0].text.strip()
    if "System Owner" in header:
        contact["role"] = "System Owner"
    elif "ISSO" in header:
        contact["role"] = "ISSO"
    elif "AO" in header or "Authorizing Official" in header:
        contact["role"] = "Authorizing Official"
    else:
        contact["role"] = header.split("(")[0].strip() if "(" in header else header
    
    for row in table.rows[1:]:  # skip header
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip()
        value = row.cells[1].text.strip()
        
        for key_prefix, field_name in field_map.items():
            if label.startswith(key_prefix):
                if value.startswith("<") or "<Enter" in value or "<555" in value:
                    value = ""
                contact[field_name] = value
                break
    
    return contact


def parse_leveraged_services(table):
    """Extract leveraged FedRAMP-authorized services."""
    services = []
    headers = [c.text.strip()[:30] for c in table.rows[0].cells]
    
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):  # skip empty rows
            continue
        
        service = {}
        for i, cell_value in enumerate(cells):
            if i < len(headers):
                if cell_value and not cell_value.startswith("<"):
                    service[headers[i]] = cell_value
        
        if service:
            services.append(service)
    
    return services


def parse_ports_protocols(table):
    """Extract services, ports, and protocols table."""
    entries = []
    headers = [c.text.strip()[:30] for c in table.rows[0].cells]
    
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue
        
        entry = {}
        for i, cell_value in enumerate(cells):
            if i < len(headers):
                if cell_value and not cell_value.startswith("<"):
                    entry[headers[i]] = cell_value
        
        if entry:
            entries.append(entry)
    
    return entries


def parse_main_ssp(docx_path):
    """Parse the main SSP template for system info, contacts, and services."""
    print(f"  Loading Main SSP: {docx_path}...")
    doc = Document(docx_path)
    print(f"  Found {len(doc.tables)} tables")
    
    result = {
        "system_information": {},
        "contacts": [],
        "leveraged_services": [],
        "ports_protocols": [],
        "separation_of_duties": [],
        "cryptographic_modules": [],    
        "information_types": []
    }
    
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        first_cell = table.rows[0].cells[0].text.strip()
        
        # System Information table
        if first_cell == "System Information":
            result["system_information"] = parse_system_information(table)
        
        # Contact tables (System Owner, ISSO, etc.)
        elif any(keyword in first_cell for keyword in ["System Owner", "ISSO", "Authorizing Official", "Point of Contact"]):
            contact = parse_contact_table(table)
            if contact.get("role"):
                result["contacts"].append(contact)
        
        # Leveraged services
        elif first_cell == "#" and len(table.rows[0].cells) > 5:
            header_text = " ".join(c.text[:20] for c in table.rows[0].cells)
            if "CSP" in header_text or "Service" in header_text:
                services = parse_leveraged_services(table)
                result["leveraged_services"].extend(services)
        
        # Ports and protocols
        elif "Port" in first_cell or "Protocol" in first_cell or "Service" in first_cell:
            if len(table.rows[0].cells) >= 3:
                entries = parse_ports_protocols(table)
                result["ports_protocols"].extend(entries)
        
        # Information types
        elif "Information Type" in first_cell and len(table.rows[0].cells) > 5:
            entries = parse_ports_protocols(table)
            if entries:
                result["information_types"] = entries
        
        # Separation of duties
        elif "Duty" in first_cell or ("Role" in first_cell and "Duties" in first_cell):
            entries = parse_ports_protocols(table)  # same row/column extraction
            result["separation_of_duties"].extend(entries)
    
    return result


# ── BUILD OUTPUT ──

def build_control_output(controls):
    """Build structured output from parsed controls."""
    output_controls = []
    
    for cid in sorted(controls.keys(), key=lambda x: (x.split("-")[0], int(re.search(r'\d+', x.split("-")[1]).group()))):
        ctrl = controls[cid]
        summary = ctrl.get("summary", {})
        impl = ctrl.get("implementation", {})
        
        filled_parts = sum(1 for v in impl.get("parts", {}).values() if v)
        total_parts = len(impl.get("parts", {}))
        
        entry = {
            "control_id": cid,
            "responsible_role": summary.get("responsible_role", ""),
            "parameters": summary.get("parameters", {}),
            "implementation_status": summary.get("implementation_status", ""),
            "control_origination": summary.get("control_origination", ""),
            "implementation_parts": impl.get("parts", {}),
            "completeness": {
                "filled_parts": filled_parts,
                "total_parts": total_parts,
                "complete": filled_parts == total_parts and total_parts > 0
            }
        }
        output_controls.append(entry)
    
    return output_controls


def build_output(main_ssp_data, controls):
    """Build the final structured output combining main SSP and Appendix A."""
    control_list = build_control_output(controls)
    
    total = len(control_list)
    complete = sum(1 for c in control_list if c["completeness"]["complete"])
    has_content = sum(1 for c in control_list if c["completeness"]["filled_parts"] > 0)
    
    output = {
        "parser_version": "2.0.0",
        "source": "FedRAMP SSP",
        "system_information": main_ssp_data.get("system_information", {}),
        "contacts": main_ssp_data.get("contacts", []),
        "leveraged_services": main_ssp_data.get("leveraged_services", []),
        "ports_protocols": main_ssp_data.get("ports_protocols", []),
        "separation_of_duties": main_ssp_data.get("separation_of_duties", []),
        "total_controls": total,
        "controls": control_list,
        "statistics": {
            "total_controls": total,
            "controls_with_content": has_content,
            "fully_complete": complete,
            "empty": total - has_content,
            "completion_rate": f"{(has_content / total * 100):.1f}%" if total > 0 else "0%",
            "has_system_info": bool(main_ssp_data.get("system_information")),
            "contacts_count": len(main_ssp_data.get("contacts", [])),
            "leveraged_services_count": len(main_ssp_data.get("leveraged_services", []))
        }
    }
    
    return output


# ── MAIN ──

def main():
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python3 ssp_parser.py <appendix-a.docx> [--output results.json]")
        print("  python3 ssp_parser.py <main-ssp.docx> --appendix-a <appendix-a.docx> [--output results.json]")
        print()
        print("Examples:")
        print("  python3 ssp_parser.py SSP-Appendix-A-Moderate.docx")
        print("  python3 ssp_parser.py FedRAMP-SSP.docx --appendix-a SSP-Appendix-A-Moderate.docx")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    appendix_a_path = None
    output_path = "ssp-parsed.json"
    
    if "--appendix-a" in sys.argv:
        idx = sys.argv.index("--appendix-a")
        if idx + 1 < len(sys.argv):
            appendix_a_path = sys.argv[idx + 1]
    
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]
    
    print("=" * 60)
    print("OpenFRAMP SSP Parser v2.0")
    print("=" * 60)
    
    main_ssp_data = {}
    controls = {}
    
    # Detect document type
    doc = Document(docx_path)
    first_tables_text = " ".join(
        doc.tables[i].rows[0].cells[0].text[:50] 
        for i in range(min(10, len(doc.tables))) 
        if doc.tables[i].rows
    )
    
    if "Control Summary Information" in first_tables_text:
        # This is Appendix A
        print("  Detected: Appendix A (control implementations)")
        controls = parse_appendix_a(docx_path)
    else:
        # This is the main SSP
        print("  Detected: Main SSP template")
        main_ssp_data = parse_main_ssp(docx_path)
    
    # If Appendix A provided separately
    if appendix_a_path:
        controls = parse_appendix_a(appendix_a_path)
    
    # Build output
    output = build_output(main_ssp_data, controls)
    
    with open(output_path, "w") as f:
        json.dump(output, f, indent=2)
    
    # Print summary
    stats = output["statistics"]
    sys_info = output.get("system_information", {})
    
    print()
    print("=" * 60)
    print("SSP Parser Results")
    print("=" * 60)
    
    if sys_info:
        print(f"System name:           {sys_info.get('cso_name', '(not set)')}")
        print(f"CSP:                   {sys_info.get('csp_name', '(not set)')}")
        print(f"Service model:         {sys_info.get('service_model', '(not set)')}")
        print(f"FIPS 199 level:        {sys_info.get('fips_199_level', '(not set)')}")
        print(f"Deployment model:      {sys_info.get('deployment_model', '(not set)')}")
        print()
    
    print(f"Contacts found:        {stats['contacts_count']}")
    print(f"Leveraged services:    {stats['leveraged_services_count']}")
    print(f"Controls found:        {stats['total_controls']}")
    print(f"With content:          {stats['controls_with_content']}")
    print(f"Fully complete:        {stats['fully_complete']}")
    print(f"Empty (template only): {stats['empty']}")
    print(f"Completion rate:       {stats['completion_rate']}")
    print(f"Output:                {output_path}")
    print("=" * 60)
    
    if output["contacts"]:
        print()
        print("Contacts:")
        for c in output["contacts"]:
            name = c.get("name", "(blank)")
            role = c.get("role", "")
            print(f"  {role}: {name}")
    
    if stats['total_controls'] > 0:
        print()
        print("Preview (first 3 controls):")
        for ctrl in output["controls"][:3]:
            parts_count = ctrl["completeness"]["total_parts"]
            filled = ctrl["completeness"]["filled_parts"]
            print(f"  {ctrl['control_id']}: {parts_count} parts ({filled} filled)")


if __name__ == "__main__":
    main()
