#!/usr/bin/env python3
"""
Unit tests for OpenFRAMP SSP Parser.

Tests three variants:
1. Appendix A only (control implementations)
2. Main SSP only (system info, contacts)
3. Combined (both documents)

Run: python3 -m pytest test_ssp_parser.py -v
  or: python3 test_ssp_parser.py
"""

import json
import os
import sys
import tempfile
import unittest
from docx import Document
from docx.shared import Inches


# ── TEST DOCUMENT BUILDERS ──

def create_test_appendix_a(path):
    """Create a minimal Appendix A with 3 controls."""
    doc = Document()
    
    # Control 1: AC-1 with summary and implementation
    # Summary table
    t1 = doc.add_table(rows=5, cols=2)
    t1.rows[0].cells[0].text = "AC-1 Control Summary Information"
    t1.rows[1].cells[0].text = "Responsible Role"
    t1.rows[1].cells[1].text = "Security Officer"
    t1.rows[2].cells[0].text = "Parameter AC-1(a):"
    t1.rows[2].cells[1].text = "annually"
    t1.rows[3].cells[0].text = "Implementation Status: Implemented"
    t1.rows[4].cells[0].text = "Control Origination: Service Provider Corporate"
    
    doc.add_paragraph("")
    
    # Implementation table
    t2 = doc.add_table(rows=4, cols=2)
    t2.rows[0].cells[0].text = "AC-1 What is the solution and how is it implemented?"
    t2.rows[1].cells[0].text = "Part a:"
    t2.rows[1].cells[1].text = "The organization develops and disseminates an access control policy."
    t2.rows[2].cells[0].text = "Part b:"
    t2.rows[2].cells[1].text = "The organization reviews and updates the access control policy annually."
    t2.rows[3].cells[0].text = "Part c:"
    t2.rows[3].cells[1].text = ""
    
    doc.add_paragraph("")
    
    # Control 2: AC-2 summary only (no implementation)
    t3 = doc.add_table(rows=3, cols=2)
    t3.rows[0].cells[0].text = "AC-2 Control Summary Information"
    t3.rows[1].cells[0].text = "Responsible Role"
    t3.rows[1].cells[1].text = "System Administrator"
    t3.rows[2].cells[0].text = "Implementation Status: Partially Implemented"
    
    doc.add_paragraph("")
    
    # Control 3: SC-28 with implementation
    t4 = doc.add_table(rows=3, cols=2)
    t4.rows[0].cells[0].text = "SC-28 Control Summary Information"
    t4.rows[1].cells[0].text = "Responsible Role"
    t4.rows[1].cells[1].text = "Cloud Engineer"
    t4.rows[2].cells[0].text = "Implementation Status: Implemented"
    
    doc.add_paragraph("")
    
    t5 = doc.add_table(rows=3, cols=2)
    t5.rows[0].cells[0].text = "SC-28 What is the solution and how is it implemented?"
    t5.rows[1].cells[0].text = "Part a:"
    t5.rows[1].cells[1].text = "All data at rest is encrypted using AES-256 with customer-managed KMS keys."
    t5.rows[2].cells[0].text = "Part b:"
    t5.rows[2].cells[1].text = "Key rotation is enabled with annual rotation schedule."
    
    doc.save(path)


def create_test_main_ssp(path):
    """Create a minimal main SSP with system info and contacts."""
    doc = Document()
    
    # Pad with empty tables to match expected positions
    for _ in range(9):
        t = doc.add_table(rows=1, cols=1)
        t.rows[0].cells[0].text = "Filler"
        doc.add_paragraph("")
    
    # Table 9: System Information
    t_sys = doc.add_table(rows=7, cols=2)
    t_sys.rows[0].cells[0].text = "System Information"
    t_sys.rows[0].cells[1].text = "System Information"
    t_sys.rows[1].cells[0].text = "CSP Name:"
    t_sys.rows[1].cells[1].text = "TestCSP Inc"
    t_sys.rows[2].cells[0].text = "CSO Name:"
    t_sys.rows[2].cells[1].text = "TestCloud Platform"
    t_sys.rows[3].cells[0].text = "Service Model: "
    t_sys.rows[3].cells[1].text = "SaaS"
    t_sys.rows[4].cells[0].text = "FIPS PUB 199 Level (SSP Appendix K):"
    t_sys.rows[4].cells[1].text = "Moderate"
    t_sys.rows[5].cells[0].text = "Deployment Model:"
    t_sys.rows[5].cells[1].text = "Government-Only Cloud"
    t_sys.rows[6].cells[0].text = "General System Description:"
    t_sys.rows[6].cells[1].text = "TestCloud is a SaaS platform for compliance management."
    
    doc.add_paragraph("")
    
    # Table 10: System Owner
    t_owner = doc.add_table(rows=7, cols=2)
    t_owner.rows[0].cells[0].text = "System Owner Information"
    t_owner.rows[0].cells[1].text = "System Owner Information"
    t_owner.rows[1].cells[0].text = "Name"
    t_owner.rows[1].cells[1].text = "Jane Smith"
    t_owner.rows[2].cells[0].text = "Title"
    t_owner.rows[2].cells[1].text = "VP of Engineering"
    t_owner.rows[3].cells[0].text = "Company / Organization"
    t_owner.rows[3].cells[1].text = "TestCSP Inc"
    t_owner.rows[4].cells[0].text = "Address"
    t_owner.rows[4].cells[1].text = "123 Cloud St, San Jose, CA 95110"
    t_owner.rows[5].cells[0].text = "Phone Number"
    t_owner.rows[5].cells[1].text = "408-555-1234"
    t_owner.rows[6].cells[0].text = "Email Address"
    t_owner.rows[6].cells[1].text = "jane.smith@testcsp.com"
    
    doc.add_paragraph("")
    
    # Filler
    t_filler = doc.add_table(rows=1, cols=1)
    t_filler.rows[0].cells[0].text = "Instructions: "
    doc.add_paragraph("")
    
    # Table 12: ISSO Contact
    t_isso = doc.add_table(rows=7, cols=2)
    t_isso.rows[0].cells[0].text = "ISSO (or Equivalent) Point of Contact"
    t_isso.rows[0].cells[1].text = "ISSO (or Equivalent) Point of Contact"
    t_isso.rows[1].cells[0].text = "Name"
    t_isso.rows[1].cells[1].text = "Bob Johnson"
    t_isso.rows[2].cells[0].text = "Title"
    t_isso.rows[2].cells[1].text = "Information Security Officer"
    t_isso.rows[3].cells[0].text = "Company / Organization"
    t_isso.rows[3].cells[1].text = "TestCSP Inc"
    t_isso.rows[4].cells[0].text = "Address"
    t_isso.rows[4].cells[1].text = "123 Cloud St, San Jose, CA 95110"
    t_isso.rows[5].cells[0].text = "Phone Number"
    t_isso.rows[5].cells[1].text = "408-555-5678"
    t_isso.rows[6].cells[0].text = "Email Address"
    t_isso.rows[6].cells[1].text = "bob.johnson@testcsp.com"
    
    doc.save(path)


def create_test_blank_template(path):
    """Create SSP with placeholder text (blank template)."""
    doc = Document()
    
    for _ in range(9):
        t = doc.add_table(rows=1, cols=1)
        t.rows[0].cells[0].text = "Filler"
        doc.add_paragraph("")
    
    t_sys = doc.add_table(rows=5, cols=2)
    t_sys.rows[0].cells[0].text = "System Information"
    t_sys.rows[0].cells[1].text = "System Information"
    t_sys.rows[1].cells[0].text = "CSP Name:"
    t_sys.rows[1].cells[1].text = "<Insert CSP Name>"
    t_sys.rows[2].cells[0].text = "CSO Name:"
    t_sys.rows[2].cells[1].text = "<Insert CSO Name>"
    t_sys.rows[3].cells[0].text = "Service Model: "
    t_sys.rows[3].cells[1].text = "<Choose one: IaaS, PaaS, SaaS>"
    t_sys.rows[4].cells[0].text = "FIPS PUB 199 Level (SSP Appendix K):"
    t_sys.rows[4].cells[1].text = "<Choose one: High, Moderate, Low>"
    
    doc.add_paragraph("")
    
    t_owner = doc.add_table(rows=4, cols=2)
    t_owner.rows[0].cells[0].text = "System Owner Information"
    t_owner.rows[0].cells[1].text = "System Owner Information"
    t_owner.rows[1].cells[0].text = "Name"
    t_owner.rows[1].cells[1].text = "<Enter Name>"
    t_owner.rows[2].cells[0].text = "Phone Number"
    t_owner.rows[2].cells[1].text = "<555-555-5555>"
    t_owner.rows[3].cells[0].text = "Email Address"
    t_owner.rows[3].cells[1].text = "<Enter Email Address>"
    
    doc.save(path)


# ── TESTS ──

class TestAppendixAParsing(unittest.TestCase):
    """Test Appendix A control extraction."""
    
    @classmethod
    def setUpClass(cls):
        cls.tmpdir = tempfile.mkdtemp()
        cls.appendix_path = os.path.join(cls.tmpdir, "test-appendix-a.docx")
        create_test_appendix_a(cls.appendix_path)
        
        # Import parser
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from ssp_parser import parse_appendix_a, build_control_output
        
        cls.controls = parse_appendix_a(cls.appendix_path)
        cls.output = build_control_output(cls.controls)
    
    def test_finds_three_controls(self):
        self.assertEqual(len(self.output), 3)
    
    def test_control_ids_correct(self):
        ids = [c["control_id"] for c in self.output]
        self.assertIn("AC-1", ids)
        self.assertIn("AC-2", ids)
        self.assertIn("SC-28", ids)
    
    def test_ac1_has_responsible_role(self):
        ac1 = next(c for c in self.output if c["control_id"] == "AC-1")
        self.assertEqual(ac1["responsible_role"], "Security Officer")
    
    def test_ac1_has_parameter(self):
        ac1 = next(c for c in self.output if c["control_id"] == "AC-1")
        self.assertIn("AC-1(a)", ac1["parameters"])
        self.assertEqual(ac1["parameters"]["AC-1(a)"], "annually")
    
    def test_ac1_has_implementation_parts(self):
        ac1 = next(c for c in self.output if c["control_id"] == "AC-1")
        self.assertIn("a", ac1["implementation_parts"])
        self.assertIn("b", ac1["implementation_parts"])
        self.assertIn("access control policy", ac1["implementation_parts"]["a"])
    
    def test_ac1_completeness(self):
        ac1 = next(c for c in self.output if c["control_id"] == "AC-1")
        self.assertEqual(ac1["completeness"]["filled_parts"], 2)  # a and b filled, c empty
        self.assertEqual(ac1["completeness"]["total_parts"], 3)
        self.assertFalse(ac1["completeness"]["complete"])  # c is empty
    
    def test_sc28_fully_complete(self):
        sc28 = next(c for c in self.output if c["control_id"] == "SC-28")
        self.assertEqual(sc28["completeness"]["filled_parts"], 2)
        self.assertEqual(sc28["completeness"]["total_parts"], 2)
        self.assertTrue(sc28["completeness"]["complete"])
    
    def test_ac2_has_no_implementation(self):
        ac2 = next(c for c in self.output if c["control_id"] == "AC-2")
        self.assertEqual(len(ac2["implementation_parts"]), 0)
    
    def test_implementation_status_extracted(self):
        ac1 = next(c for c in self.output if c["control_id"] == "AC-1")
        self.assertIn("Implemented", ac1["implementation_status"])
    
    def test_control_origination_extracted(self):
        ac1 = next(c for c in self.output if c["control_id"] == "AC-1")
        self.assertIn("Service Provider", ac1["control_origination"])


class TestMainSSPParsing(unittest.TestCase):
    """Test main SSP system info and contact extraction."""
    
    @classmethod
    def setUpClass(cls):
        cls.tmpdir = tempfile.mkdtemp()
        cls.ssp_path = os.path.join(cls.tmpdir, "test-main-ssp.docx")
        create_test_main_ssp(cls.ssp_path)
        
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from ssp_parser import parse_main_ssp
        
        cls.result = parse_main_ssp(cls.ssp_path)
    
    def test_system_info_extracted(self):
        info = self.result["system_information"]
        self.assertEqual(info["csp_name"], "TestCSP Inc")
        self.assertEqual(info["cso_name"], "TestCloud Platform")
        self.assertEqual(info["service_model"], "SaaS")
        self.assertEqual(info["fips_199_level"], "Moderate")
        self.assertEqual(info["deployment_model"], "Government-Only Cloud")
    
    def test_system_description_extracted(self):
        info = self.result["system_information"]
        self.assertIn("compliance management", info["description"])
    
    def test_contacts_found(self):
        self.assertEqual(len(self.result["contacts"]), 2)
    
    def test_system_owner_contact(self):
        owner = next(c for c in self.result["contacts"] if c["role"] == "System Owner")
        self.assertEqual(owner["name"], "Jane Smith")
        self.assertEqual(owner["title"], "VP of Engineering")
        self.assertEqual(owner["organization"], "TestCSP Inc")
        self.assertEqual(owner["email"], "jane.smith@testcsp.com")
        self.assertEqual(owner["phone"], "408-555-1234")
    
    def test_isso_contact(self):
        isso = next(c for c in self.result["contacts"] if c["role"] == "ISSO")
        self.assertEqual(isso["name"], "Bob Johnson")
        self.assertEqual(isso["title"], "Information Security Officer")
        self.assertEqual(isso["email"], "bob.johnson@testcsp.com")


class TestBlankTemplate(unittest.TestCase):
    """Test that blank template placeholders are stripped."""
    
    @classmethod
    def setUpClass(cls):
        cls.tmpdir = tempfile.mkdtemp()
        cls.path = os.path.join(cls.tmpdir, "test-blank.docx")
        create_test_blank_template(cls.path)
        
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from ssp_parser import parse_main_ssp
        
        cls.result = parse_main_ssp(cls.path)
    
    def test_placeholder_text_stripped(self):
        info = self.result["system_information"]
        self.assertEqual(info.get("csp_name", ""), "")
        self.assertEqual(info.get("cso_name", ""), "")
        self.assertEqual(info.get("service_model", ""), "")
    
    def test_placeholder_contacts_stripped(self):
        for contact in self.result["contacts"]:
            self.assertEqual(contact.get("name", ""), "")
            self.assertEqual(contact.get("phone", ""), "")
            self.assertEqual(contact.get("email", ""), "")


class TestCombinedPipeline(unittest.TestCase):
    """Test the full pipeline with both documents."""
    
    @classmethod
    def setUpClass(cls):
        cls.tmpdir = tempfile.mkdtemp()
        cls.ssp_path = os.path.join(cls.tmpdir, "test-main.docx")
        cls.appendix_path = os.path.join(cls.tmpdir, "test-appendix.docx")
        create_test_main_ssp(cls.ssp_path)
        create_test_appendix_a(cls.appendix_path)
        
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from ssp_parser import parse_main_ssp, parse_appendix_a, build_output
        
        main_data = parse_main_ssp(cls.ssp_path)
        controls = parse_appendix_a(cls.appendix_path)
        cls.output = build_output(main_data, controls)
    
    def test_combined_has_system_info(self):
        self.assertTrue(self.output["statistics"]["has_system_info"])
        self.assertEqual(self.output["system_information"]["csp_name"], "TestCSP Inc")
    
    def test_combined_has_contacts(self):
        self.assertEqual(self.output["statistics"]["contacts_count"], 2)
    
    def test_combined_has_controls(self):
        self.assertEqual(self.output["statistics"]["total_controls"], 3)
    
    def test_combined_completion_stats(self):
        stats = self.output["statistics"]
        self.assertEqual(stats["fully_complete"], 1)  # SC-28
        self.assertEqual(stats["controls_with_content"], 2)  # AC-1 and SC-28
        self.assertEqual(stats["empty"], 1)  # AC-2
    
    def test_parser_version(self):
        self.assertEqual(self.output["parser_version"], "2.0.0")


class TestOSCALGeneration(unittest.TestCase):
    """Test OSCAL SSP JSON generation."""
    
    @classmethod
    def setUpClass(cls):
        cls.tmpdir = tempfile.mkdtemp()
        cls.ssp_path = os.path.join(cls.tmpdir, "test-main.docx")
        cls.appendix_path = os.path.join(cls.tmpdir, "test-appendix.docx")
        create_test_main_ssp(cls.ssp_path)
        create_test_appendix_a(cls.appendix_path)
        
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from ssp_parser import parse_main_ssp, parse_appendix_a, build_output
        from ssp_to_oscal import build_oscal_ssp
        
        main_data = parse_main_ssp(cls.ssp_path)
        controls = parse_appendix_a(cls.appendix_path)
        combined = build_output(main_data, controls)
        cls.oscal = build_oscal_ssp(combined)
    
    def test_oscal_has_system_security_plan(self):
        self.assertIn("system-security-plan", self.oscal)
    
    def test_oscal_version(self):
        self.assertEqual(self.oscal["system-security-plan"]["metadata"]["oscal-version"], "1.1.2")
    
    def test_oscal_system_name(self):
        name = self.oscal["system-security-plan"]["system-characteristics"]["system-name"]
        self.assertEqual(name, "TestCloud Platform")
    
    def test_oscal_security_level(self):
        level = self.oscal["system-security-plan"]["system-characteristics"]["security-sensitivity-level"]
        self.assertEqual(level, "moderate")
    
    def test_oscal_has_parties(self):
        parties = self.oscal["system-security-plan"]["metadata"]["parties"]
        self.assertGreaterEqual(len(parties), 3)  # 2 contacts + 1 org
    
    def test_oscal_has_implemented_requirements(self):
        reqs = self.oscal["system-security-plan"]["control-implementation"]["implemented-requirements"]
        self.assertEqual(len(reqs), 3)
    
    def test_oscal_control_ids_lowercase(self):
        reqs = self.oscal["system-security-plan"]["control-implementation"]["implemented-requirements"]
        for req in reqs:
            self.assertEqual(req["control-id"], req["control-id"].lower())
    
    def test_oscal_ac1_has_statements(self):
        reqs = self.oscal["system-security-plan"]["control-implementation"]["implemented-requirements"]
        ac1 = next(r for r in reqs if r["control-id"] == "ac-1")
        self.assertIn("statements", ac1)
        self.assertGreater(len(ac1["statements"]), 0)
    
    def test_oscal_has_import_profile(self):
        profile = self.oscal["system-security-plan"]["import-profile"]["href"]
        self.assertIn("MODERATE", profile)
    
    def test_oscal_system_characteristics_props(self):
        chars = self.oscal["system-security-plan"]["system-characteristics"]
        props = chars.get("props", [])
        prop_names = [p["name"] for p in props]
        self.assertIn("cloud-service-model", prop_names)
        self.assertIn("cloud-deployment-model", prop_names)


# ── RUN ──

if __name__ == "__main__":
    # Run with verbose output
    unittest.main(verbosity=2)
